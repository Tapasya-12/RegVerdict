"""
Thin HTTP wrapper around mcp_server/tools.check_compliance, for the React
frontend (frontend/) to call directly over fetch() — the MCP server itself speaks stdio/MCP
protocol, not plain HTTP, so this is a separate small FastAPI process rather
than a change to mcp_server/server.py.

Run with: uvicorn server:app --reload --port 8000 (from this directory).
"""

import sys
from contextlib import asynccontextmanager
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "mcp_server"))
sys.path.insert(0, str(Path(__file__).resolve().parent))

# auth.py reads JWT_SECRET_KEY from the environment at import time, so the
# .env it lives in (mcp_server/.env — reused here rather than adding a
# second env file, since llm_client.py already loads it the same way) must
# be loaded before `import auth` runs, not left to depend on import order.
from dotenv import load_dotenv  # noqa: E402

load_dotenv(Path(__file__).resolve().parent.parent / "mcp_server" / ".env")

import io  # noqa: E402
from datetime import datetime, timezone  # noqa: E402

import groq  # noqa: E402
import jwt  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402
from fastapi import Depends, FastAPI, HTTPException, Request  # noqa: E402
from fastapi.middleware.cors import CORSMiddleware  # noqa: E402
from fastapi.responses import JSONResponse, StreamingResponse  # noqa: E402
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm  # noqa: E402
from pydantic import BaseModel  # noqa: E402
from slowapi import Limiter  # noqa: E402
from slowapi.errors import RateLimitExceeded  # noqa: E402
from slowapi.util import get_remote_address  # noqa: E402

import audit_log_sqlite  # noqa: E402
import auth  # noqa: E402
import recent_queries  # noqa: E402
import tools  # noqa: E402


@asynccontextmanager
async def lifespan(app: FastAPI):
    auth.init_db()
    yield


app = FastAPI(title="RegVerdict API", lifespan=lifespan)

# Vite picks whichever port is free (5173, 5174, ...) — match any localhost
# port instead of hardcoding one, since the frontend is a separate dev process.
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r"http://(localhost|127\.0\.0\.1):\d+",
    allow_methods=["*"],
    allow_headers=["*"],
    # Content-Disposition carries export_report's real filename — browsers
    # always get it, but the Fetch API hides all but a small "safe" set of
    # response headers from JS on cross-origin requests unless the server
    # explicitly opts it in here.
    expose_headers=["Content-Disposition"],
)


oauth2_scheme = OAuth2PasswordBearer(tokenUrl="api/auth/login")


def get_current_user(token: str = Depends(oauth2_scheme)) -> str:
    try:
        return auth.decode_access_token(token)
    except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
        raise HTTPException(status_code=401, detail="Invalid or expired session")


def rate_limit_key(request: Request) -> str:
    # Keyed by the authenticated user, not IP — multiple users can share a
    # network (office wifi, NAT), which would make IP-based limiting either
    # too strict (one heavy user blocks everyone behind the same router) or
    # too loose (nothing stops one user cycling networks). The token is
    # already validated by get_current_user's own Depends() on these routes
    # by the time slowapi's wrapper runs, so decoding failures here should
    # be unreachable — the IP fallback only exists so a key_func call can
    # never itself throw and take down the request.
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        token = auth_header[len("Bearer "):]
        try:
            return auth.decode_access_token(token)
        except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
            pass
    return get_remote_address(request)


limiter = Limiter(key_func=rate_limit_key)
app.state.limiter = limiter


@app.exception_handler(RateLimitExceeded)
def rate_limit_exceeded_handler(request: Request, exc: RateLimitExceeded) -> JSONResponse:
    return JSONResponse(
        status_code=429,
        content={"detail": "Too many requests — please wait a moment before submitting another check."},
    )


def get_current_user_id(username: str = Depends(get_current_user)) -> int:
    # audit_log_sqlite/recent_queries key everything off the real integer
    # user_id (foreign key to users.id), never the username string — this
    # is the one place that resolves the validated token back to that id.
    user = auth.get_user_by_username(username)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid or expired session")
    return user["id"]


class SignupRequest(BaseModel):
    username: str
    email: str
    password: str


class RecentQueryCreate(BaseModel):
    full_query: str
    display_title: str | None = None


class RecentQueryRename(BaseModel):
    display_title: str


class CheckComplianceRequest(BaseModel):
    policy_text: str


class SimulatePolicyChangeRequest(BaseModel):
    original_policy: str
    proposed_change: str


class CompareJurisdictionsRequest(BaseModel):
    policy_text: str
    regulators: list[str]


@app.get("/api/health")
def health() -> dict:
    # Also pre-warms tools.get_retriever()'s singleton (embedding model +
    # BM25 index build) so the first real check_compliance call isn't the
    # one paying that cold-start cost. Deliberately unprotected — useful for
    # infra checks that shouldn't need a session.
    retriever = tools.get_retriever()
    return {"status": "ok", "indexed_chunks": len(retriever.payload_by_id)}


@app.post("/api/auth/signup")
def signup(req: SignupRequest) -> dict:
    try:
        return auth.create_user(req.username, req.email, req.password)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/auth/login")
def login(form_data: OAuth2PasswordRequestForm = Depends()) -> dict:
    user = auth.authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(status_code=401, detail="Incorrect username or password")
    return {"access_token": auth.create_access_token(user["username"]), "token_type": "bearer"}


def _run_or_groq_503(fn, *args):
    # The 20/minute app-level limiter above caps request COUNT per user, but
    # Groq's own quota is tokens-per-minute across the whole shared account —
    # a burst of legitimate, allowed requests can still exhaust it (confirmed
    # via a 25-concurrent-request test: request count was capped correctly,
    # but one of the 20 allowed requests still hit groq.RateLimitError and
    # surfaced as a raw unhandled 500). This turns that into the same kind
    # of clean, expected response as our own rate limit, distinguishable by
    # message from the app-level 429 above.
    try:
        return fn(*args)
    except groq.GroqError as e:
        raise HTTPException(
            status_code=503,
            detail="The compliance engine's shared quota is temporarily exhausted — please wait a few seconds and try again.",
        ) from e
    except RuntimeError as e:
        # rag/retriever.py's embedded-mode Qdrant client can't safely serve
        # true concurrent access (see the comment in dense_search()) — under
        # heavy simultaneous load this is the specific error it raises.
        # Only local dev without Docker hits this; a real Qdrant server
        # (QDRANT_URL) doesn't have the limitation. Narrowed to this exact
        # message so an unrelated RuntimeError elsewhere in the pipeline
        # still surfaces as a real 500, not silently masked as "busy."
        if "already accessed by another instance" in str(e):
            raise HTTPException(
                status_code=503,
                detail="The compliance engine is handling another request right now — please try again in a moment.",
            ) from e
        raise


def _verdict_str(v):
    # verdict_output.model_dump() (mode="python", the default) leaves
    # "verdict" as a raw Verdict enum member — fine for FastAPI's own JSON
    # response encoding, but str(Verdict.NON_COMPLIANT) would render as
    # "Verdict.NON_COMPLIANT" instead of "Non-Compliant" anywhere else
    # (audit log rows, summary strings), so callers normalize through here.
    return v.value if hasattr(v, "value") else v


@app.post("/api/check_compliance")
@limiter.limit("20/minute")
def check_compliance(
    request: Request, req: CheckComplianceRequest, user_id: int = Depends(get_current_user_id)
) -> dict:
    if not req.policy_text or not req.policy_text.strip():
        raise HTTPException(status_code=400, detail="policy_text must be a non-empty string")

    result = _run_or_groq_503(tools.check_compliance, req.policy_text)

    source_clause = result.get("source_clause") or {}
    audit_log_sqlite.append_log(
        auth.DB_PATH,
        user_id,
        tool="check_compliance",
        policy_text=req.policy_text,
        verdict=_verdict_str(result.get("verdict")),
        confidence=result.get("confidence"),
        document_name=source_clause.get("document") if source_clause else None,
        clause_number=source_clause.get("clause_number") if source_clause else None,
        grounding_verified=result.get("grounding_verified"),
    )

    # This endpoint is only ever called from the Workspace tool (Policy Diff
    # and Compare Jurisdictions call tools.check_compliance() directly as a
    # Python function via their own endpoints, not this route) — so every
    # hit here is a real Workspace submission that belongs in the sidebar.
    recent_queries.create_recent_query(auth.DB_PATH, user_id, req.policy_text)

    return result


@app.post("/api/simulate_policy_change")
@limiter.limit("20/minute")
def simulate_policy_change(
    request: Request, req: SimulatePolicyChangeRequest, user_id: int = Depends(get_current_user_id)
) -> dict:
    if not req.original_policy or not req.original_policy.strip():
        raise HTTPException(status_code=400, detail="original_policy must be a non-empty string")
    if not req.proposed_change or not req.proposed_change.strip():
        raise HTTPException(status_code=400, detail="proposed_change must be a non-empty string")

    result = _run_or_groq_503(tools.simulate_policy_change, req.original_policy, req.proposed_change)

    original_verdict = _verdict_str(result.get("original_verdict", {}).get("verdict"))
    proposed_verdict = _verdict_str(result.get("proposed_verdict", {}).get("verdict"))
    audit_log_sqlite.append_log(
        auth.DB_PATH,
        user_id,
        tool="simulate_policy_change",
        policy_text=req.original_policy,
        proposed_change=req.proposed_change,
        verdict=f"{original_verdict} → {proposed_verdict}",
        status_flipped=result.get("status_flipped"),
    )

    return result


@app.post("/api/compare_jurisdictions")
@limiter.limit("20/minute")
def compare_jurisdictions(
    request: Request, req: CompareJurisdictionsRequest, user_id: int = Depends(get_current_user_id)
) -> dict:
    if not req.policy_text or not req.policy_text.strip():
        raise HTTPException(status_code=400, detail="policy_text must be a non-empty string")
    if not req.regulators:
        raise HTTPException(status_code=400, detail="regulators must be a non-empty list")

    result = _run_or_groq_503(tools.compare_jurisdictions, req.policy_text, req.regulators)

    # No single clause/confidence to log here — each regulator in the
    # comparison has its own, so the audit entry carries a per-regulator
    # verdict summary instead (mirrors simulate_policy_change's
    # "before → after" summary string for the same reason).
    results_by_regulator = result.get("results_by_regulator", {})
    verdict_summary = " · ".join(
        f"{regulator}: {_verdict_str(r.get('verdict'))}"
        for regulator, r in results_by_regulator.items()
    )
    audit_log_sqlite.append_log(
        auth.DB_PATH,
        user_id,
        tool="compare_jurisdictions",
        policy_text=req.policy_text,
        verdict=verdict_summary,
    )

    return result


@app.get("/api/documents")
def documents(current_user: str = Depends(get_current_user)) -> list[str]:
    retriever = tools.get_retriever()
    return sorted({
        p.get("document_name") for p in retriever.payload_by_id.values()
        if p.get("document_name")
    })


@app.get("/api/clause_graph/{document_name}")
def clause_graph(document_name: str, current_user: str = Depends(get_current_user)) -> dict:
    # Read-only structural query, not a compliance verdict — no audit_log entry.
    result = tools.get_clause_graph(document_name)
    if "error" in result:
        raise HTTPException(status_code=404, detail=result["error"])
    return result


# ---------------------------------------------------------------------------
# fetch_regulation / generate_compliance_report / detect_regulatory_conflicts /
# get_regulation_history: these 4 MCP tools (mcp_server/tools.py) were only
# ever wired into the MCP protocol server (mcp_server/server.py) for stdio
# MCP clients — checked git history, they were never exposed here for the
# React frontend, so there's nothing to "restore" so much as add for the
# first time. None of the 5 built frontend tools call these yet, so — like
# clause_graph/documents above — no audit_log entry or recent_queries side
# effect; they're plain protected reads, same pattern as the rest of this
# file's read-only endpoints.
# ---------------------------------------------------------------------------

@app.get("/api/fetch_regulation")
def fetch_regulation(topic: str, top_k: int = 5, current_user: str = Depends(get_current_user)) -> dict:
    if not topic or not topic.strip():
        raise HTTPException(status_code=400, detail="topic must be a non-empty string")
    return tools.fetch_regulation(topic, top_k)


@app.post("/api/generate_compliance_report")
def generate_compliance_report(
    req: CheckComplianceRequest, current_user: str = Depends(get_current_user)
) -> dict:
    if not req.policy_text or not req.policy_text.strip():
        raise HTTPException(status_code=400, detail="policy_text must be a non-empty string")
    return _run_or_groq_503(tools.generate_compliance_report, req.policy_text)


@app.post("/api/detect_regulatory_conflicts")
def detect_regulatory_conflicts(
    req: CheckComplianceRequest, current_user: str = Depends(get_current_user)
) -> dict:
    if not req.policy_text or not req.policy_text.strip():
        raise HTTPException(status_code=400, detail="policy_text must be a non-empty string")
    return _run_or_groq_503(tools.detect_regulatory_conflicts, req.policy_text)


@app.get("/api/regulation_history/{clause_id}")
def get_regulation_history(clause_id: str, current_user: str = Depends(get_current_user)) -> dict:
    return tools.get_regulation_history(clause_id)


def _build_report_docx(report: dict, policy_text: str, current_user: str) -> tuple[io.BytesIO, str]:
    # Pure rendering — no HTTP/auth concerns — so it's testable directly
    # against a real captured report dict without needing a live server
    # request (useful when Groq's own quota is what's under test).
    verdict = report.get("verdict", {})

    doc = Document()
    doc.add_heading("RegVerdict Compliance Report", level=1)

    doc.add_heading("Policy Received", level=2)
    doc.add_paragraph(report.get("policy_received", policy_text))

    doc.add_heading("Relevant Regulations Found", level=2)
    relevant_regs = report.get("relevant_regulations_found", [])
    if relevant_regs:
        for reg in relevant_regs:
            doc.add_paragraph(
                f"{reg.get('document_name')} §{reg.get('clause_number')} ({reg.get('regulator')})",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No relevant regulations found in the indexed corpus.")

    doc.add_heading("Verdict", level=2)
    doc.add_paragraph(verdict.get("policy_summary", ""))

    verdict_p = doc.add_paragraph()
    verdict_p.add_run("Verdict: ").bold = True
    verdict_p.add_run(_verdict_str(verdict.get("verdict")) or "—")

    confidence_p = doc.add_paragraph()
    confidence_p.add_run("Confidence: ").bold = True
    confidence = verdict.get("confidence")
    confidence_p.add_run(f"{confidence:.2f}" if isinstance(confidence, (int, float)) else "—")

    doc.add_paragraph("Evidence Quote:").runs[0].bold = True
    quote_p = doc.add_paragraph()
    quote_p.paragraph_format.left_indent = Inches(0.5)
    quote_text = verdict.get("evidence_quote")
    quote_run = quote_p.add_run(f'"{quote_text}"' if quote_text else "— no verbatim span available —")
    quote_run.italic = True

    doc.add_heading("Reasoning", level=3)
    doc.add_paragraph(verdict.get("reasoning", "—"))

    doc.add_heading("Recommended Action", level=3)
    doc.add_paragraph(verdict.get("recommended_action", "—"))

    generated_at = datetime.now(timezone.utc)
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        f"Generated by RegVerdict on {generated_at.strftime('%Y-%m-%d %H:%M UTC')} for {current_user}"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"regverdict_report_{generated_at.strftime('%Y%m%d_%H%M%S')}.docx"
    return buffer, filename


@app.post("/api/export_report")
def export_report(req: CheckComplianceRequest, current_user: str = Depends(get_current_user)) -> StreamingResponse:
    # Re-runs generate_compliance_report() fresh rather than accepting a
    # verdict payload from the client — a client-supplied verdict could be
    # tampered with before export, and this is the same tool the frontend
    # already displayed the verdict from, so the .docx matches what the
    # user saw. Not rate-limited (unlike check_compliance/etc.) per this
    # task's explicit scope, though it does call Groq the same as those —
    # flagged, not silently decided.
    if not req.policy_text or not req.policy_text.strip():
        raise HTTPException(status_code=400, detail="policy_text must be a non-empty string")

    report = _run_or_groq_503(tools.generate_compliance_report, req.policy_text)
    buffer, filename = _build_report_docx(report, req.policy_text, current_user)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/audit_trail")
def audit_trail(
    verdict: str | None = None,
    start_date: str | None = None,
    end_date: str | None = None,
    search: str | None = None,
    current_user: str = Depends(get_current_user),
) -> list[dict]:
    # Deliberately not scoped to current_user — an audit trail's whole point
    # is visibility across the firm's compliance activity, not a private
    # history. (Contrast with recent_queries below, which is per-user by
    # design and enforces it at the storage layer, not just the query.)
    return audit_log_sqlite.read_log(
        auth.DB_PATH, verdict=verdict, start_date=start_date, end_date=end_date, search=search
    )


@app.post("/api/recent_queries")
def create_recent_query_endpoint(
    req: RecentQueryCreate, user_id: int = Depends(get_current_user_id)
) -> dict:
    return recent_queries.create_recent_query(auth.DB_PATH, user_id, req.full_query, req.display_title)


@app.get("/api/recent_queries")
def list_recent_queries_endpoint(user_id: int = Depends(get_current_user_id)) -> list[dict]:
    return recent_queries.list_recent_queries(auth.DB_PATH, user_id)


@app.patch("/api/recent_queries/{query_id}/pin")
def pin_recent_query_endpoint(query_id: int, user_id: int = Depends(get_current_user_id)) -> dict:
    try:
        return recent_queries.toggle_pin(auth.DB_PATH, query_id, user_id)
    except ValueError as e:
        # Same message whether the id doesn't exist at all or belongs to
        # someone else — the caller shouldn't be able to tell those apart.
        raise HTTPException(status_code=404, detail=str(e))


@app.patch("/api/recent_queries/{query_id}")
def rename_recent_query_endpoint(
    query_id: int, req: RecentQueryRename, user_id: int = Depends(get_current_user_id)
) -> dict:
    try:
        return recent_queries.rename_query(auth.DB_PATH, query_id, user_id, req.display_title)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.delete("/api/recent_queries/{query_id}")
def delete_recent_query_endpoint(query_id: int, user_id: int = Depends(get_current_user_id)) -> dict:
    deleted = recent_queries.delete_query(auth.DB_PATH, query_id, user_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Query not found or does not belong to this user.")
    return {"deleted": True}
